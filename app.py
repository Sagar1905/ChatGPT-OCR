import os
from flask import Flask, request, render_template, send_file, redirect, url_for
from werkzeug.utils import secure_filename
from docx import Document
from openai import OpenAI
from PIL import Image
import base64
from dotenv import load_dotenv
import re
import html
import cv2
import numpy as np
from datetime import datetime

# Load environment variables from .env file
load_dotenv()

app = Flask(__name__, template_folder='templates')
app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Initialize OpenAI client (will be None if API key is not set)
api_key = os.getenv('OPENAI_API_KEY')
client = OpenAI(api_key=api_key) if api_key else None


def preprocess_image(image_path: str) -> str:
    """Preprocess image to improve OCR accuracy."""
    try:
        print(f"Starting preprocessing for: {image_path}")
        
        # Read the image
        img = cv2.imread(image_path)
        if img is None:
            # If OpenCV can't read it, try with PIL and convert
            from PIL import Image as PILImage
            pil_img = PILImage.open(image_path)
            img = cv2.cvtColor(np.array(pil_img), cv2.COLOR_RGB2BGR)
        
        original_height, original_width = img.shape[:2]
        print(f"Original image size: {original_width}x{original_height}")
        
        # Convert to grayscale
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        
        # Light noise reduction (more conservative)
        denoised = cv2.medianBlur(gray, 3)
        
        # Conservative contrast enhancement
        clahe = cv2.createCLAHE(clipLimit=1.5, tileGridSize=(8,8))
        enhanced = clahe.apply(denoised)
        
        # Only upscale if image is quite small
        height, width = enhanced.shape[:2]
        if height < 800 or width < 800:
            scale_factor = min(1200/height, 1200/width, 2.0)  # More conservative scaling
            new_width = int(width * scale_factor)
            new_height = int(height * scale_factor)
            enhanced = cv2.resize(enhanced, (new_width, new_height), interpolation=cv2.INTER_CUBIC)
            print(f"Scaled image to: {new_width}x{new_height}")
        
        # Save the preprocessed image (keep as grayscale, don't apply binary threshold)
        processed_path = image_path.replace('.', '_processed.')
        success = cv2.imwrite(processed_path, enhanced)
        
        if not success:
            print("Failed to save preprocessed image, using original")
            return image_path
        
        print(f"Preprocessing completed: {processed_path}")
        return processed_path
        
    except Exception as e:
        print(f"Preprocessing error: {e}, using original image")
        return image_path


def detect_document_type(image_path: str) -> str:
    """Detect the type of document to use appropriate extraction strategy."""
    img = cv2.imread(image_path, cv2.IMREAD_GRAYSCALE)
    if img is None:
        return "general"
    
    height, width = img.shape
    aspect_ratio = width / height
    
    # Detect lines and text density
    edges = cv2.Canny(img, 50, 150, apertureSize=3)
    lines = cv2.HoughLinesP(edges, 1, np.pi/180, threshold=100, minLineLength=100, maxLineGap=10)
    
    horizontal_lines = 0
    vertical_lines = 0
    
    if lines is not None:
        for line in lines:
            x1, y1, x2, y2 = line[0]
            angle = np.arctan2(y2 - y1, x2 - x1) * 180 / np.pi
            if abs(angle) < 10 or abs(angle) > 170:
                horizontal_lines += 1
            elif abs(abs(angle) - 90) < 10:
                vertical_lines += 1
    
    # Classification logic
    if horizontal_lines > 10 and vertical_lines > 5:
        return "table"
    elif horizontal_lines > 15:
        return "form"
    elif aspect_ratio > 1.5:
        return "receipt"
    else:
        return "document"


def get_enhanced_prompt(doc_type: str) -> str:
    """Get specialized prompts based on document type."""
    
    if doc_type == "table":
        return "What text and data do you see in this table? Please list everything you can read."
    
    elif doc_type == "form":
        return "What text do you see in this form? Please read all the fields and values."
    
    elif doc_type == "receipt":
        return "What text do you see on this receipt? Please read all the items and amounts."
    
    else:  # document
        return "What text do you see in this document? Please read everything you can see."


def extract_text_with_layout(image_path: str, use_preprocessing: bool = True, high_accuracy: bool = False) -> str:
    """Use OpenAI to extract text from an image with enhanced accuracy.

    Args:
        image_path: Path to the image file
        use_preprocessing: Whether to preprocess the image for better OCR
        high_accuracy: Whether to use multi-pass extraction for highest accuracy
    """
    if client is None:
        raise ValueError("OpenAI API key not set. Please set the OPENAI_API_KEY environment variable.")
    
    print(f"Starting extraction for: {image_path}")
    print(f"Settings: preprocessing={use_preprocessing}, high_accuracy={high_accuracy}")
    
    # Multiple fallback strategies with different approaches
    strategies = [
        ("enhanced", use_preprocessing, True),   # Enhanced with preprocessing
        ("simple", False, False),               # Simple without preprocessing
        ("basic", False, False),                # Basic fallback
        ("minimal", False, False),              # Minimal prompt as last resort
        ("ultra_basic", False, False),          # Ultra basic - original image, basic prompt
        ("ocr_focused", False, False),          # OCR-focused approach
    ]
    
    for strategy_name, use_preprocess, use_enhanced_prompt in strategies:
        try:
            print(f"Trying {strategy_name} strategy...")
            
            # Step 1: Detect document type for specialized handling (only for enhanced)
            if use_enhanced_prompt:
                doc_type = detect_document_type(image_path)
                print(f"Detected document type: {doc_type}")
            else:
                doc_type = "document"
            
            # Step 2: Preprocess image if enabled
            if use_preprocess:
                try:
                    processed_image_path = preprocess_image(image_path)
                    print(f"Image preprocessed: {processed_image_path}")
                except Exception as prep_error:
                    print(f"Preprocessing failed: {prep_error}, using original image")
                    processed_image_path = image_path
            else:
                processed_image_path = image_path
            
            # Step 3: Get prompt based on strategy
            if use_enhanced_prompt:
                prompt = get_enhanced_prompt(doc_type)
            elif strategy_name == "minimal":
                prompt = get_minimal_prompt()
            elif strategy_name == "ultra_basic":
                prompt = get_ultra_basic_prompt()
            elif strategy_name == "ocr_focused":
                prompt = get_ocr_focused_prompt()
            else:
                prompt = get_simple_prompt()
            
            # Step 4: Perform extraction
            with open(processed_image_path, 'rb') as img_file:
                base64_image = base64.b64encode(img_file.read()).decode('utf-8')
                
                print(f"Image size: {len(base64_image)} characters")
                
                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "text",
                                    "text": prompt,
                                },
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:image/jpeg;base64,{base64_image}"
                                    }
                                }
                            ]
                        }
                    ],
                    max_tokens=4096,
                    temperature=0.1,
                )
            
            extracted_text = response.choices[0].message.content
            print(f"Extracted text length: {len(extracted_text)} characters")
            
            # Log the actual response for debugging
            if len(extracted_text) < 200:
                print(f"Short response content: {extracted_text[:100]}...")
            else:
                print(f"Long response preview: {extracted_text[:50]}...")
            
            # Log exact response for debugging problematic cases
            print(f"Full response (first 200 chars): {repr(extracted_text[:200])}")
            
            # Improved failure detection - be more selective about what we consider a failure
            if extracted_text and len(extracted_text.strip()) > 5:
                # Check for definitive refusal messages (but not if they also contain extracted text)
                definitive_failures = [
                    "I'm unable to extract text",
                    "I cannot extract text", 
                    "I'm sorry, I can't extract text",
                    "I cannot see any text",
                    "Unable to process",
                    "I don't see any text",
                    "I'm unable to transcribe",
                    "I cannot transcribe",
                    "I'm unable to read the text",
                    "I cannot read the text",
                    "I'm unable to make out",
                    "I cannot make out",
                    "I don't see any readable text",
                    "I'm not able to extract",
                    "I'm not able to read"
                ]
                
                # Check if the response is a clear refusal
                is_refusal = False
                extracted_lower = extracted_text.strip().lower()
                
                # First check: direct refusal patterns (more flexible matching)
                for failure_msg in definitive_failures:
                    if failure_msg.lower() in extracted_lower:
                        is_refusal = True
                        print(f"Detected refusal pattern: '{failure_msg}' in response")
                        break
                
                # Second check: common refusal sentence starters
                refusal_starters = [
                    "i'm unable to extract text from",
                    "i cannot extract text from", 
                    "i'm sorry, i can't extract",
                    "i'm unable to transcribe the text",
                    "i cannot transcribe the text",
                    "i'm unable to read",
                    "i cannot read",
                    "i don't see any text",
                    "i'm not able to extract"
                ]
                
                if not is_refusal:
                    for starter in refusal_starters:
                        if extracted_lower.startswith(starter):
                            is_refusal = True
                            print(f"Detected refusal starter: '{starter}'")
                            break
                
                # Third check: if response is short and contains key refusal indicators
                if not is_refusal and len(extracted_text.strip()) < 300:
                    # Check if it's primarily a refusal/apology message
                    refusal_indicators = ['unable', 'cannot', "can't", 'sorry', "don't see", 'not able', 'provide a clearer', 'happy to help']
                    indicator_count = sum(1 for indicator in refusal_indicators if indicator in extracted_lower)
                    
                    # If it has multiple refusal indicators and no actual content, it's likely a refusal
                    if indicator_count >= 2:
                        # Check for content indicators
                        content_words = extracted_lower.split()
                        actual_content_words = [word for word in content_words if word not in 
                                              ['i', 'am', 'unable', 'to', 'extract', 'text', 'from', 'the', 'image', 'you', 'provided', 
                                               'if', 'can', 'provide', 'a', 'clearer', 'or', 'describe', 'content', 'id', 'be', 'happy', 
                                               'help', 'further', 'sorry', 'cant', 'assistance', 'with', 'something', 'else', 'feel', 
                                               'free', 'ask', 'might', 'able']]
                        
                        if len(actual_content_words) < 3:  # Very few non-refusal words
                            is_refusal = True
                            print(f"Detected refusal based on content analysis. Indicator count: {indicator_count}, Content words: {actual_content_words}")
                
                if is_refusal:
                    print(f"Strategy {strategy_name} failed with refusal message")
                    continue
                
                # If we have any text that looks like content, consider it a success
                if len(extracted_text.strip()) > 10:
                    print(f"Strategy {strategy_name} succeeded!")
                    
                    # Step 5: High accuracy mode - second pass for verification (only if requested and first strategy)
                    if high_accuracy and strategy_name == "enhanced" and len(extracted_text) > 50:
                        try:
                            print("Performing high-accuracy second pass...")
                            
                            verification_prompt = f"""Please review and improve this text extraction. Extract all visible text from the image, focusing on accuracy:

Previous extraction attempt:
{extracted_text[:500]}...

Please provide a complete, accurate extraction of all text visible in the image:"""
                            
                            with open(processed_image_path, 'rb') as img_file:
                                base64_image = base64.b64encode(img_file.read()).decode('utf-8')
                                
                                verification_response = client.chat.completions.create(
                                    model="gpt-4o",
                                    messages=[
                                        {
                                            "role": "user",
                                            "content": [
                                                {
                                                    "type": "text",
                                                    "text": verification_prompt,
                                                },
                                                {
                                                    "type": "image_url",
                                                    "image_url": {
                                                        "url": f"data:image/jpeg;base64,{base64_image}"
                                                    }
                                                }
                                            ]
                                        }
                                    ],
                                    max_tokens=4096,
                                    temperature=0.4,
                                )
                            
                            verified_text = verification_response.choices[0].message.content
                            if verified_text and len(verified_text.strip()) > len(extracted_text.strip()) * 0.8:
                                extracted_text = verified_text
                                print("High-accuracy pass completed successfully")
                        except Exception as verify_error:
                            print(f"Verification pass failed: {verify_error}, using original extraction")
                    
                    # Step 6: Cleanup - remove processed image if it was created
                    if use_preprocess and processed_image_path != image_path:
                        try:
                            os.remove(processed_image_path)
                        except:
                            pass  # Ignore cleanup errors
                    
                    return extracted_text
            else:
                print(f"Strategy {strategy_name} returned insufficient text")
                continue
                
        except Exception as e:
            print(f"Strategy {strategy_name} failed with error: {str(e)}")
            continue
    
    # If all strategies failed, return a helpful error message
    print("All extraction strategies failed")
    return "Unable to extract text from this image. The image may be too complex, corrupted, or contain content that cannot be processed. Please try with a clearer, higher-resolution image."


def get_simple_prompt() -> str:
    """Get a simple, reliable prompt for text extraction."""
    return "What text do you see in this image? Please read everything you can see."


def get_minimal_prompt() -> str:
    """Get a minimal prompt as a last resort."""
    return "What text do you see?"


def get_ultra_basic_prompt() -> str:
    """Get an ultra-basic prompt as a last resort."""
    return "What text is in this image?"


def get_ocr_focused_prompt() -> str:
    """Get an OCR-focused prompt as a last resort."""
    return "What text and numbers can you read in this image? Please list everything you see."


def save_to_docx(markdown_text: str, output_path: str):
    """Save markdown text as a properly formatted docx document."""
    doc = Document()
    
    # First, unescape HTML entities
    clean_text = html.unescape(markdown_text)
    
    lines = clean_text.splitlines()
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines but add paragraph breaks
        if not line:
            if i > 0:  # Don't add paragraph at the very beginning
                doc.add_paragraph()
            i += 1
            continue
        
        # Handle headers
        if line.startswith('#'):
            level = 0
            while level < len(line) and line[level] == '#':
                level += 1
            
            header_text = line[level:].strip()
            if level == 1:
                doc.add_heading(header_text, level=1)
            elif level == 2:
                doc.add_heading(header_text, level=2)
            elif level == 3:
                doc.add_heading(header_text, level=3)
            else:
                doc.add_heading(header_text, level=4)
        
        # Handle bullet points
        elif line.startswith(('- ', '* ', '• ')):
            bullet_text = line[2:].strip()
            formatted_text = format_inline_markdown(bullet_text)
            para = doc.add_paragraph(style='List Bullet')
            add_formatted_text(para, formatted_text)
        
        # Handle numbered lists
        elif re.match(r'^\d+\.\s', line):
            list_text = re.sub(r'^\d+\.\s', '', line).strip()
            formatted_text = format_inline_markdown(list_text)
            para = doc.add_paragraph(style='List Number')
            add_formatted_text(para, formatted_text)
        
        # Handle table rows (simple approach)
        elif '|' in line and line.count('|') >= 2:
            # Simple table handling - just add as formatted paragraph
            formatted_text = format_inline_markdown(line.replace('|', ' | '))
            para = doc.add_paragraph()
            add_formatted_text(para, formatted_text)
        
        # Handle regular paragraphs
        else:
            formatted_text = format_inline_markdown(line)
            para = doc.add_paragraph()
            add_formatted_text(para, formatted_text)
        
        i += 1
    
    doc.save(output_path)


def format_inline_markdown(text: str) -> list:
    """Parse inline markdown formatting and return list of (text, formatting) tuples."""
    result = []
    
    # Split by bold and italic markers
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|__[^_]+__|_[^_]+_)', text)
    
    for part in parts:
        if not part:
            continue
        
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            result.append((part[2:-2], {'bold': True}))
        elif part.startswith('__') and part.endswith('__'):
            # Bold text (alternative)
            result.append((part[2:-2], {'bold': True}))
        elif part.startswith('*') and part.endswith('*'):
            # Italic text
            result.append((part[1:-1], {'italic': True}))
        elif part.startswith('_') and part.endswith('_'):
            # Italic text (alternative)
            result.append((part[1:-1], {'italic': True}))
        else:
            # Regular text
            result.append((part, {}))
    
    return result


def add_formatted_text(paragraph, formatted_parts):
    """Add formatted text parts to a paragraph."""
    for text, formatting in formatted_parts:
        run = paragraph.add_run(text)
        if formatting.get('bold'):
            run.bold = True
        if formatting.get('italic'):
            run.italic = True


@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        # Check if API key is set
        if client is None:
            error_message = "OpenAI API key not set. Please set the OPENAI_API_KEY environment variable."
            return render_template('index.html', error=error_message)
        
        files = request.files.getlist('documents')
        doc_paths = []
        error_message = None
        
        try:
            for file in files:
                if not file.filename:
                    continue
                filename = secure_filename(file.filename)
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(file_path)

                # Get accuracy settings from form
                use_preprocessing = request.form.get('preprocessing', 'on') == 'on'
                high_accuracy = request.form.get('high_accuracy', 'off') == 'on'
                
                print(f"Processing {filename} with preprocessing={use_preprocessing}, high_accuracy={high_accuracy}")
                
                # Extract text and layout using OpenAI with enhanced accuracy
                markdown_text = extract_text_with_layout(
                    file_path, 
                    use_preprocessing=use_preprocessing, 
                    high_accuracy=high_accuracy
                )
                
                # Generate unique filename with timestamp
                timestamp = datetime.now().strftime("%d%m%Y%H%M%S")
                base_name = os.path.splitext(filename)[0]
                output_docx = os.path.join(app.config['UPLOAD_FOLDER'], f"{timestamp}_{base_name}.docx")
                
                save_to_docx(markdown_text, output_docx)
                doc_paths.append(output_docx)

            # Display only first document text for preview
            preview_text = ""
            if doc_paths:
                with open(doc_paths[0], 'rb') as doc_file:
                    # Convert docx text to plain text for preview
                    from docx import Document as DocReader
                    doc_reader = DocReader(doc_file)
                    preview_text = "\n".join([p.text for p in doc_reader.paragraphs])
            return render_template('index.html', docs=doc_paths, preview=preview_text)
        
        except Exception as e:
            print(f"Full error details: {str(e)}")
            import traceback
            traceback.print_exc()
            error_message = f"Error processing documents: {str(e)}"
            return render_template('index.html', error=error_message)
            
    # Show API key status on GET request
    api_key_status = "✅ OpenAI API key is set" if client else "❌ OpenAI API key not set"
    return render_template('index.html', docs=None, api_key_status=api_key_status)


@app.route('/download/<path:filename>')
def download(filename):
    return send_file(filename, as_attachment=True)


if __name__ == '__main__':
    app.run(debug=True)
